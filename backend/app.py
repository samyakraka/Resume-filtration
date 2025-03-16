import os
import re
import json
import PyPDF2
import docx
import numpy as np
from datetime import datetime
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity

class ATSScorer:
    def __init__(self):
        self.weights = {
            'skills': 0.25,
            'experience': 0.20,
            'education': 0.15,
            'formatting': 0.10,
            'achievements': 0.00,
            'projects': 0.25,
            'cgpa': 0.05
        }
        
        self.json_storage = "resume_database.json"

    def read_and_save_resume(self, file_path):
        """Read resume content and save/update JSON database with improved text processing"""
        file_ext = os.path.splitext(file_path)[1].lower()
        content = ""
        
        try:
            if file_ext == '.pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        # Extract and preprocess text from each page
                        page_text = page.extract_text()
                        processed_text = self.preprocess_text(page_text)
                        content += processed_text + "\n"
            elif file_ext in ['.docx', '.doc']:
                doc = docx.Document(file_path)
                # Process each paragraph
                paragraphs = [self.preprocess_text(paragraph.text) for paragraph in doc.paragraphs]
                content = '\n'.join(paragraphs)
            else:
                raise ValueError("Unsupported file format")
            
            # Additional cleaning for the entire content
            content = self.clean_resume_content(content)
            
            # Save to database
            self.save_to_database(file_path, content)
            
        except Exception as e:
            raise Exception(f"Error processing resume file: {str(e)}")
            
        return content
    def preprocess_text(self, text):
        """Clean and normalize text for better matching"""
        # Remove extra whitespace and normalize spaces
        text = re.sub(r'\s+', ' ', text)
        
        # Add space before capital letters that might indicate new words
        text = re.sub(r'(?<!^)(?<![\s.])([A-Z][a-z])', r' \1', text)
        
        # Handle common abbreviations and prevent unwanted splits
        text = re.sub(r'(?<=B)\.(?=Tech|E|A|S)', '', text, flags=re.IGNORECASE)
        text = re.sub(r'(?<=M)\.(?=Tech|E|A|S)', '', text, flags=re.IGNORECASE)
        
        # Normalize education-related terms
        text = re.sub(r'B[\s.]*Tech', 'Bachelor of Technology', text, flags=re.IGNORECASE)
        text = re.sub(r'M[\s.]*Tech', 'Master of Technology', text, flags=re.IGNORECASE)
        text = re.sub(r'B[\s.]*E', 'Bachelor of Engineering', text, flags=re.IGNORECASE)
        text = re.sub(r'M[\s.]*E', 'Master of Engineering', text, flags=re.IGNORECASE)
        
        return text.strip()
    def clean_resume_content(self, content):
        """Advanced cleaning of resume content"""
        # Replace multiple newlines with single newline
        content = re.sub(r'\n\s*\n', '\n', content)
        
        # Add space after common section headers
        content = re.sub(r'(Education|Experience|Skills|Projects):', r'\1: ', content, flags=re.IGNORECASE)
        
        # Fix common formatting issues
        content = re.sub(r'([a-z])([A-Z])', r'\1 \2', content)  # Add space between camelCase
        content = re.sub(r'([A-Z]{2,})([a-z])', r'\1 \2', content)  # Add space after UPPERCASE words
        
        # Normalize bullets and lists
        content = re.sub(r'[•●∙○◦⚬⦁]', '- ', content)
        
        # Remove unnecessary symbols while preserving important ones
        content = re.sub(r'[^\w\s\-.,()%$:/@+]', ' ', content)
        
        # Normalize spaces
        content = ' '.join(content.split())
        
        return content    
    def save_to_database(self, file_path, content):
        """Save processed resume to database"""
        if os.path.exists(self.json_storage):
            with open(self.json_storage, 'r', encoding='utf-8') as file:
                database = json.load(file)
        else:
            database = {"resumes": []}
        
        filename = os.path.basename(file_path)
        new_entry = {
            "filename": filename,
            "last_updated": datetime.now().isoformat(),
            "content": content
        }
        
        # Update or add entry
        for idx, entry in enumerate(database["resumes"]):
            if entry["filename"] == filename:
                database["resumes"][idx] = new_entry
                break
        else:
            database["resumes"].append(new_entry)
        
        with open(self.json_storage, 'w', encoding='utf-8') as file:
            json.dump(database, file, indent=4)
    def read_and_save_resume(self, file_path):
        """Read resume content and save/update JSON database"""
        file_ext = os.path.splitext(file_path)[1].lower()
        content = ""
        
        try:
            if file_ext == '.pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        content += page.extract_text()
            elif file_ext in ['.docx', '.doc']:
                doc = docx.Document(file_path)
                content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            else:
                raise ValueError("Unsupported file format")
            
            # Load existing database or create new one
            if os.path.exists(self.json_storage):
                with open(self.json_storage, 'r', encoding='utf-8') as file:
                    database = json.load(file)
            else:
                database = {"resumes": []}
            
            # Update or add new entry
            filename = os.path.basename(file_path)
            new_entry = {
                "filename": filename,
                "last_updated": datetime.now().isoformat(),
                "content": content
            }
            
            # Check if file already exists in database
            for idx, entry in enumerate(database["resumes"]):
                if entry["filename"] == filename:
                    database["resumes"][idx] = new_entry
                    break
            else:
                database["resumes"].append(new_entry)
            
            # Save updated database
            with open(self.json_storage, 'w', encoding='utf-8') as file:
                json.dump(database, file, indent=4)
                
        except Exception as e:
            raise Exception(f"Error processing resume file: {str(e)}")
            
        return content

    def parse_job_description(self, job_desc):
        """Parse job description into structured requirements"""
        requirements = {
            'skills': [],
            'experience': {'years': 0, 'type': ''},
            'education': '',
            'cgpa': {'min': 0.0, 'scale': 4.0}  # Default to 4.0 scale
        }
        
        # Extract skills
        skills_match = re.search(r'Skills:\s*(.*?)\.', job_desc)
        if skills_match:
            requirements['skills'] = [s.strip() for s in skills_match.group(1).split(',')]
        
        # Extract experience
        exp_match = re.search(r'Experience:\s*(\d+)\+?\s*years?\s*in\s*(.*?)\.', job_desc)
        if exp_match:
            requirements['experience'] = {
                'years': int(exp_match.group(1)),
                'type': exp_match.group(2).strip()
            }
        
        # Extract education
        edu_match = re.search(r'Education:\s*(.*?)\.', job_desc)
        if edu_match:
            requirements['education'] = edu_match.group(1).strip()
        
        # Extract CGPA requirement
        cgpa_match = re.search(r'CGPA:\s*(\d+(?:\.\d+)?)\s*(?:\/\s*(\d+(?:\.\d+)?))?', job_desc)
        if cgpa_match:
            requirements['cgpa']['min'] = float(cgpa_match.group(1))
            if cgpa_match.group(2):
                requirements['cgpa']['scale'] = float(cgpa_match.group(2))
            
        return requirements
    def generate_skill_variations(self, skill):
        """Generate common variations of skill names"""
        variations = [skill]
        
        # Remove spaces
        variations.append(skill.replace(' ', ''))
        
        # Handle common abbreviations
        if skill.lower() in ['javascript', 'typescript']:
            variations.extend([f"{skill[:2]}", skill.replace('script', '')])
        elif 'python' in skill.lower():
            variations.extend(['py', 'python2', 'python3'])
        elif 'amazon web services' in skill.lower():
            variations.extend(['aws', 'amazon aws'])
        
        # Add lowercase and uppercase variations
        variations.extend([v.lower() for v in variations])
        variations.extend([v.upper() for v in variations])
        
        return list(set(variations))
    
    def calculate_skills_score(self, resume_text, required_skills):
        """Calculate skills match score with fuzzy matching"""
        found_skills = 0
        total_skills = len(required_skills)
        matched_skills = []
        missing_skills = []
        
        # Normalize resume text for better matching
        normalized_text = resume_text.lower()
        
        for skill in required_skills:
            skill_variations = self.generate_skill_variations(skill)
            found = False
            
            for variation in skill_variations:
                if re.search(rf'\b{re.escape(variation)}\b', normalized_text, re.IGNORECASE):
                    found = True
                    matched_skills.append(skill)  # Store original skill name
                    break
            
            if found:
                found_skills += 1
            else:
                missing_skills.append(skill)
        
        score = (found_skills / total_skills) * 100 if total_skills > 0 else 0
        return score, matched_skills, missing_skills


    def calculate_experience_score(self, resume_text, required_exp):
        """Calculate experience match score"""
        exp_patterns = [
            rf'\b(\d+)\+?\s*years?\s*(?:of)?\s*experience\s*(?:in)?\s*{required_exp["type"]}',
            rf'\b(\d+)\+?\s*years?\s*(?:in)?\s*{required_exp["type"]}'
        ]
        
        found_years = 0
        for pattern in exp_patterns:
            match = re.search(pattern, resume_text, re.IGNORECASE)
            if match:
                found_years = int(match.group(1))
                break
        
        if found_years >= required_exp['years']:
            score = 100
        else:
            score = (found_years / required_exp['years']) * 100 if required_exp['years'] > 0 else 0
            
        return min(score, 100), found_years

    def calculate_achievements_score(self, resume_text):
        """Calculate score based on quantifiable achievements with dynamic patterns"""
        achievement_patterns = [
            # Numeric achievements
            r'\b(\d+(?:\.\d+)?%?\s*(?:increase|decrease|improvement|growth|reduction))',
            r'\$\s*\d+(?:,\d+)*(?:\.\d+)?(?:\s*[kmbt])?(?:\s*(?:USD|EUR|GBP))?',
            r'\b(\d+)(?:,\d+)*\s*(?:users|customers|clients|employees|team members|projects)',
            
            # Time-based achievements
            r'\b(?:reduced|decreased|cut)\s*time\s*by\s*\d+(?:\.\d+)?%',
            r'\b(?:completed|delivered)\s*\d+\s*(?:months|weeks|days)\s*ahead',
            
            # Efficiency improvements
            r'\b(?:improved|increased|enhanced)\s*efficiency\s*by\s*\d+(?:\.\d+)?%',
            
            # Revenue/cost impact
            r'\b(?:generated|earned|saved|reduced costs?)\s*(?:by)?\s*\$\s*\d+(?:,\d+)*(?:\.\d+)?[kmbt]?',
            
            # Team/project scale
            r'\b(?:led|managed|coordinated)\s*(?:a\s*)?team\s*of\s*\d+',
            r'\b(?:handled|managed|oversaw)\s*\d+\s*(?:concurrent|simultaneous)?\s*projects',
            
            # Performance metrics
            r'\b(?:achieved|maintained)\s*\d+(?:\.\d+)?%\s*(?:uptime|accuracy|satisfaction|performance)'
        ]
        
        achievements_found = []
        for pattern in achievement_patterns:
            matches = re.finditer(pattern, resume_text, re.IGNORECASE)
            achievements_found.extend([m.group(0) for m in matches])
        
        # Remove duplicates while preserving order
        achievements_found = list(dict.fromkeys(achievements_found))
        
        # Score calculation with diminishing returns
        num_achievements = len(achievements_found)
        if num_achievements == 0:
            score = 0
        else:
            # Logarithmic scoring to handle diminishing returns
            score = min(100, 25 * np.log2(num_achievements + 1))
            
        return score, achievements_found

    def calculate_projects_score(self, resume_text):
        """Calculate score based on projects and their complexity"""
        project_patterns = [
            # Project headers
            r'(?:^|\n)(?:projects?|implementations?|applications?)[:]\s*(.*?)(?:\n|$)',
            
            # Project descriptions
            r'\b(?:developed|created|implemented|designed|built)\s+(?:a|an)?\s*(?:\w+\s*){2,7}(?:system|application|platform|solution|website|tool)',
            
            # Technical project indicators
            r'\b(?:using|with|via|through)\s+(?:\w+,\s*)*(?:and\s+)?\w+\s*(?:framework|library|technology|stack)',
            
            # Project outcomes
            r'(?:resulting in|leading to|achieved)\s+(?:\w+\s*){3,10}'
        ]
        
        projects_found = []
        for pattern in project_patterns:
            matches = re.finditer(pattern, resume_text, re.IGNORECASE)
            projects_found.extend([m.group(0).strip() for m in matches])
        
        # Remove duplicates and very short matches
        projects_found = list(set([p for p in projects_found if len(p) > 20]))
        
        # Score calculation
        num_projects = len(projects_found)
        if num_projects == 0:
            score = 0
        else:
            # Logarithmic scoring with cap at 5 projects
            score = min(100, 40 * np.log2(num_projects + 1))
            
        return score, projects_found

    def calculate_cgpa_score(self, resume_text, required_cgpa):
        """Calculate score based on CGPA/percentage match"""
        # Common CGPA and percentage patterns
        patterns = [
            # CGPA patterns
            r'(?:cgpa|gpa)(?:\s*:)?\s*(\d+(?:\.\d+)?)\s*(?:/\s*(\d+(?:\.\d+)?)|$)',
            # Percentage patterns
            r'(?:percentage|%)(?:\s*:)?\s*(\d+(?:\.\d+)?)\s*%',
            # Direct number patterns
            r'(?:with|secured|obtained)(?:\s+an?)?\s*(\d+(?:\.\d+)?)\s*(?:cgpa|gpa|%)'
        ]
        
        found_value = None
        found_scale = None
        
        for pattern in patterns:
            match = re.search(pattern, resume_text, re.IGNORECASE)
            if match:
                found_value = float(match.group(1))
                if pattern.startswith('(?:cgpa|gpa)') and match.group(2):
                    found_scale = float(match.group(2))
                elif pattern.find('%') != -1:
                    found_value /= 100.0
                    found_scale = 1.0
                break
        
        if not found_value:
            return 0, None
        
        # Normalize to required scale
        if found_scale and found_scale != required_cgpa['scale']:
            normalized_value = (found_value / found_scale) * required_cgpa['scale']
        else:
            normalized_value = found_value
        
        # Calculate score
        if normalized_value >= required_cgpa['min']:
            score = 100
        else:
            score = (normalized_value / required_cgpa['min']) * 100
            
        return min(score, 100), found_value

    def calculate_education_score(self, resume_text, required_education):
        """Enhanced education match score calculation"""
        # Normalize texts
        resume_text = self.preprocess_text(resume_text.lower())
        required_education = self.preprocess_text(required_education.lower())
        
        degree_mappings = {
            "bachelor": {
                "variations": ["bachelor", "bachelors", "b.s.", "b.a.", "bs", "ba", "b.tech", "b.e.", "btech", "be",
                             "bachelor of", "undergraduate"],
                "level": 1
            },
            "master": {
                "variations": ["master", "masters", "m.s.", "m.a.", "ms", "ma", "m.tech", "m.e.", "mtech", "me",
                             "master of", "graduate"],
                "level": 2
            },
            "phd": {
                "variations": ["phd", "ph.d.", "doctorate", "doctoral", "doctor of philosophy"],
                "level": 3
            }
        }
        
        # Find required degree level
        required_level = 0
        for degree, info in degree_mappings.items():
            if any(var in required_education for var in info["variations"]):
                required_level = info["level"]
                break
        
        # Find highest degree in resume
        highest_level = 0
        for degree, info in degree_mappings.items():
            if any(var in resume_text for var in info["variations"]):
                highest_level = max(highest_level, info["level"])
        
        # Extract field of study from requirement
        field_pattern = r'(?:in|of)\s+([\w\s]+)(?:\s|$)'
        field_match = re.search(field_pattern, required_education)
        required_field = field_match.group(1) if field_match else ""
        
        # Check if field matches
        field_score = 0
        if required_field and required_field.lower() in resume_text:
            field_score = 20
        
        # Calculate base score based on degree level
        if highest_level >= required_level:
            base_score = 80
        elif highest_level == required_level - 1:
            base_score = 60
        else:
            base_score = 40
        
        return min(base_score + field_score, 100)

    def calculate_formatting_score(self, resume_text):
        """Calculate formatting score based on structure and readability"""
        score = 100
        
        required_sections = ['experience', 'education', 'skills', 'projects']
        for section in required_sections:
            if not re.search(rf'\b{section}\b', resume_text, re.IGNORECASE):
                score -= 15
        
        if len(resume_text.split('\n\n')) < 3:
            score -= 20
        
        word_count = len(resume_text.split())
        if word_count > 1000:
            score -= 15
        elif word_count < 200:
            score -= 15
            
        return max(score, 0)
        
    def evaluate_resume(self, resume_path, job_description):
        """Enhanced main evaluation function"""
        # Read and preprocess resume
        resume_text = self.read_and_save_resume(resume_path)
        normalized_resume = self.preprocess_text(resume_text)
        
        # Parse requirements
        requirements = self.parse_job_description(job_description)
        
        # Calculate scores with normalized text
        scores_data = {
            'skills': {
                'score': 0,
                'matched': [],
                'missing': []
            },
            'experience': {
                'score': 0,
                'years': 0,
                'required': requirements['experience']
            },
            'education': {
                'score': 0,
                'required': requirements['education']
            },
            'formatting': {
                'score': 0
            },
            'achievements': {
                'score': 0,
                'found': []
            },
            'projects': {
                'score': 0,
                'found': []
            },
            'cgpa': {
                'score': 0,
                'found': None,
                'required': requirements['cgpa']
            }
        }
        
        # Calculate individual scores
        scores_data['skills']['score'], scores_data['skills']['matched'], scores_data['skills']['missing'] = \
            self.calculate_skills_score(normalized_resume, requirements['skills'])
            
        scores_data['experience']['score'], scores_data['experience']['years'] = \
            self.calculate_experience_score(normalized_resume, requirements['experience'])
            
        scores_data['education']['score'] = \
            self.calculate_education_score(normalized_resume, requirements['education'])
            
        scores_data['formatting']['score'] = self.calculate_formatting_score(normalized_resume)
        
        scores_data['achievements']['score'], scores_data['achievements']['found'] = \
            self.calculate_achievements_score(normalized_resume)
            
        scores_data['projects']['score'], scores_data['projects']['found'] = \
            self.calculate_projects_score(normalized_resume)
            
        scores_data['cgpa']['score'], scores_data['cgpa']['found'] = \
            self.calculate_cgpa_score(normalized_resume, requirements['cgpa'])
        
        # Calculate overall score
        overall_score = sum(
            scores_data[category]['score'] * self.weights[category]
            for category in self.weights.keys()
        )
        
        # Generate suggestions
        suggestions = self.generate_suggestions(scores_data)
        
        return {
            'scores': scores_data,
            'overall_score': overall_score,
            'suggestions': suggestions
        }
    def generate_suggestions(self, scores_data):
        """Generate improvement suggestions based on scores"""
        suggestions = []
        
        # Skills suggestions
        if scores_data['skills']['missing']:
            suggestions.append(f"Add missing key skills: {', '.join(scores_data['skills']['missing'])}")
        
        # Experience suggestions
        if scores_data['experience']['score'] < 100:
            suggestions.append("Clarify your years of experience more explicitly")
        
        # Education suggestions
        if scores_data['education']['score'] < 80:
            suggestions.append("Ensure your education qualifications match the requirements and are clearly stated")
        
        # Achievements suggestions
        if scores_data['achievements']['score'] < 70:
            suggestions.append("Add more quantifiable achievements with specific metrics, numbers, and results")
        
        # Projects suggestions
        if scores_data['projects']['score'] < 70:
            suggestions.append("Include more detailed project descriptions with technologies used and outcomes achieved")
        
        # CGPA suggestions
        if scores_data['cgpa']['score'] < 80:
            suggestions.append("Clearly mention your CGPA/percentage in the education section")
        
        # Formatting suggestions
        if scores_data['formatting']['score'] < 80:
            suggestions.append("Improve resume structure with clear section headers and consistent formatting")
        
        return suggestions

    def format_report(self, report):
        """Format the evaluation report as a string"""
        scores = report['scores']
        
        output = []
        
        # Skills section
        skills_text = f"Matched {len(scores['skills']['matched'])}/{len(scores['skills']['matched']) + len(scores['skills']['missing'])} key skills: {', '.join(scores['skills']['matched'])}"
        output.append(f"Skills Assessment\n{skills_text}\nScore: {scores['skills']['score']:.0f}%")
        
        # Experience section
        exp_text = f"Required: {scores['experience']['required']['years']}+ years in {scores['experience']['required']['type']}\nFound: {scores['experience']['years']} years"
        output.append(f"\nExperience Assessment\n{exp_text}\nScore: {scores['experience']['score']:.0f}%")
        
        # Education section
        edu_text = f"Required: {scores['education']['required']}"
        output.append(f"\nEducation Assessment\n{edu_text}\nScore: {scores['education']['score']:.0f}%")
        
        # CGPA section
        cgpa_text = f"Required: {scores['cgpa']['required']['min']}/{scores['cgpa']['required']['scale']}"
        if scores['cgpa']['found']:
            cgpa_text += f"\nFound: {scores['cgpa']['found']}"
        output.append(f"\nCGPA Assessment\n{cgpa_text}\nScore: {scores['cgpa']['score']:.0f}%")
        
        # Achievements section
        achievements_text = f"Found {len(scores['achievements']['found'])} quantifiable achievements:"
        if scores['achievements']['found']:
            achievements_text += "\n- " + "\n- ".join(scores['achievements']['found'])
        output.append(f"\nAchievements Assessment\n{achievements_text}\nScore: {scores['achievements']['score']:.0f}%")
        
        # Projects section
        projects_text = f"Found {len(scores['projects']['found'])} relevant projects:"
        if scores['projects']['found']:
            projects_text += "\n- " + "\n- ".join(scores['projects']['found'])
        output.append(f"\nProjects Assessment\n{projects_text}\nScore: {scores['projects']['score']:.0f}%")
        
        # Formatting section
        output.append(f"\nFormatting Assessment\nScore: {scores['formatting']['score']:.0f}%")
        
        # Overall score
        output.append(f"\nOverall ATS Score: {report['overall_score']:.0f}%")
        
        # Suggestions
        if report['suggestions']:
            output.append("\nSuggested Improvements:")
            for suggestion in report['suggestions']:
                output.append(f"- {suggestion}")
        
        return "\n".join(output)

# Example usage
if __name__ == "__main__":
    # Example job description
    job_description = """
    Skills: Python, React, Nextjs.
    Experience: 0+ years in software development.
    Education: Bachelor of Technology.
    Additional Requirements: Badminton, Cricket.
    """
    
    # Initialize scorer
    scorer = ATSScorer()
    
    # Evaluate resume
    resume_path = "upload/BZ.pdf"  # Replace with actual resume path
    report = scorer.evaluate_resume(resume_path, job_description)
    
    # Print formatted report
    print(scorer.format_report(report))